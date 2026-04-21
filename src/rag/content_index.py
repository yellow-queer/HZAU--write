"""
项目内容知识库 RAG 索引模块

负责：
1. 读取多种格式的内容文件（txt, md, pdf, docx, xlsx, jpg 等）
2. 为内容创建 LlamaIndex 向量索引
3. 提供内容检索接口
4. 支持多模态输入（图像通过 Qwen-VL 转换为文字描述）
"""

import os
import time
from typing import List, Dict, Any, Optional
from pathlib import Path

from llama_index.core import (
    VectorStoreIndex,
    StorageContext,
    Settings,
    Document
)
from llama_index.core.node_parser import SentenceSplitter
from llama_index.vector_stores.chroma import ChromaVectorStore
import chromadb


class ContentIndex:
    """项目内容知识库索引"""
    
    SUPPORTED_EXTENSIONS = {
        '.txt': 'text',
        '.md': 'text',
        '.pdf': 'pdf',
        '.docx': 'word',
        '.xlsx': 'excel',
        '.xls': 'excel',
        '.jpg': 'image',
        '.jpeg': 'image',
        '.png': 'image',
    }
    
    def __init__(self, content_dir: str, persist_dir: str = ".chroma_content"):
        """
        初始化内容索引
        
        Args:
            content_dir: 内容文件夹路径
            persist_dir: ChromaDB 持久化目录
        """
        self.content_dir = Path(content_dir)
        self.persist_dir = persist_dir
        self.index: Optional[VectorStoreIndex] = None
        self.query_engine = None
        
    def build_index(self, llm=None, embed_model=None, enable_image_processing: bool = True):
        """
        构建内容向量索引
        
        Args:
            llm: LLM 模型（可选）
            embed_model: 嵌入模型（可选）
            enable_image_processing: 是否启用图像处理（需要 Qwen-VL API）
        """
        if embed_model:
            Settings.embed_model = embed_model
        else:
            try:
                from llama_index.embeddings.huggingface import HuggingFaceEmbedding
                Settings.embed_model = HuggingFaceEmbedding(
                    model_name="BAAI/bge-small-zh-v1.5",
                    cache_folder=None,
                )
                print("已加载 HuggingFace 嵌入模型：bge-small-zh-v1.5")
            except ImportError:
                try:
                    from llama_index.embeddings.langchain import LangchainEmbedding
                    from langchain_community.embeddings import HuggingFaceEmbeddings
                    Settings.embed_model = LangchainEmbedding(
                        HuggingFaceEmbeddings(model_name="BAAI/bge-small-zh-v1.5")
                    )
                    print("已加载 LangChain HuggingFace 嵌入模型：bge-small-zh-v1.5")
                except ImportError:
                    print("警告：无法加载 HuggingFace 嵌入模型，将使用默认配置")
        
        if not self.content_dir.exists():
            print(f"警告：内容文件夹不存在：{self.content_dir}")
            print(f"  请创建文件夹并放入相关内容文件")
            return
        
        documents = []
        stats = {
            'text': 0,
            'pdf': 0,
            'word': 0,
            'excel': 0,
            'image': 0,
            'failed': 0
        }
        
        print(f"\n{'='*60}")
        print(f"正在扫描内容文件夹：{self.content_dir}")
        print(f"{'='*60}")
        
        for file_path in self.content_dir.rglob("*"):
            if not file_path.is_file():
                continue
            
            suffix = file_path.suffix.lower()
            if suffix not in self.SUPPORTED_EXTENSIONS:
                continue
            
            file_type = self.SUPPORTED_EXTENSIONS[suffix]
            
            try:
                if file_type == 'text':
                    doc = self._read_text_file(file_path)
                    if doc:
                        documents.append(doc)
                        stats['text'] += 1
                        print(f"  ✓ 文本文件: {file_path.name}")
                
                elif file_type == 'pdf':
                    doc = self._read_pdf_file(file_path)
                    if doc:
                        documents.append(doc)
                        stats['pdf'] += 1
                        print(f"  ✓ PDF 文件: {file_path.name}")
                
                elif file_type == 'word':
                    doc = self._read_docx_file(file_path)
                    if doc:
                        documents.append(doc)
                        stats['word'] += 1
                        print(f"  ✓ Word 文件: {file_path.name}")
                
                elif file_type == 'excel':
                    doc = self._read_excel_file(file_path)
                    if doc:
                        documents.append(doc)
                        stats['excel'] += 1
                        print(f"  ✓ Excel 文件: {file_path.name}")
                
                elif file_type == 'image':
                    if enable_image_processing:
                        doc = self._process_image_file(file_path)
                        if doc:
                            documents.append(doc)
                            stats['image'] += 1
                            print(f"  ✓ 图像文件: {file_path.name}")
                    else:
                        print(f"  ⊘ 跳过图像文件（图像处理已禁用）: {file_path.name}")
            
            except Exception as e:
                print(f"  ✗ 读取文件失败 {file_path.name}: {e}")
                stats['failed'] += 1
        
        print(f"\n{'='*60}")
        print(f"文件扫描统计：")
        print(f"  - 文本文件: {stats['text']} 个")
        print(f"  - PDF 文件: {stats['pdf']} 个")
        print(f"  - Word 文件: {stats['word']} 个")
        print(f"  - Excel 文件: {stats['excel']} 个")
        print(f"  - 图像文件: {stats['image']} 个")
        print(f"  - 失败: {stats['failed']} 个")
        print(f"  - 总计: {len(documents)} 个文档")
        print(f"{'='*60}\n")
        
        if not documents:
            print(f"警告：未在 {self.content_dir} 中找到任何支持的文件")
            return
        
        print(f"开始构建内容索引，共 {len(documents)} 个文档...")
        start_time = time.time()
        
        text_splitter = SentenceSplitter(chunk_size=1024, chunk_overlap=128)
        
        print("正在初始化向量存储...")
        db_init_start = time.time()
        db = chromadb.PersistentClient(path=self.persist_dir)
        chroma_collection = db.get_or_create_collection("content")
        vector_store = ChromaVectorStore(chroma_collection=chroma_collection)
        storage_context = StorageContext.from_defaults(vector_store=vector_store)
        print(f"✓ 向量存储初始化完成 (耗时：{time.time() - db_init_start:.2f}秒)")
        
        print(f"正在处理 {len(documents)} 个文档...")
        batch_size = 5
        all_nodes = []
        
        for i in range(0, len(documents), batch_size):
            batch_docs = documents[i:i + batch_size]
            batch_num = (i // batch_size) + 1
            total_batches = (len(documents) + batch_size - 1) // batch_size
            
            print(f"  处理批次 {batch_num}/{total_batches} (文档 {i+1}-{min(i+batch_size, len(documents))})...")
            batch_start = time.time()
            
            try:
                batch_index = VectorStoreIndex.from_documents(
                    batch_docs,
                    storage_context=storage_context,
                    transformations=[text_splitter],
                    show_progress=True
                )
                
                if hasattr(batch_index, '_docstore'):
                    nodes = list(batch_index._docstore.docs.values())
                    all_nodes.extend(nodes)
                
                batch_time = time.time() - batch_start
                print(f"  ✓ 批次 {batch_num} 完成 (耗时：{batch_time:.2f}秒)")
                
            except Exception as e:
                print(f"  ⚠ 批次 {batch_num} 处理失败：{e}")
                import traceback
                traceback.print_exc()
                continue
        
        print(f"正在创建统一索引...")
        index_start = time.time()
        
        try:
            self.index = VectorStoreIndex(
                nodes=all_nodes if all_nodes else None,
                storage_context=storage_context,
                embed_model=Settings.embed_model
            )
            
            print(f"✓ 索引创建完成 (耗时：{time.time() - index_start:.2f}秒)")
            
            try:
                self.query_engine = self.index.as_query_engine(
                    similarity_top_k=5,
                    include_metadata=True
                )
            except Exception as e2:
                print(f"警告：查询引擎创建失败，使用简化配置：{e2}")
                from llama_index.core.query_engine import RetrieverQueryEngine
                retriever = self.index.as_retriever(similarity_top_k=5)
                self.query_engine = RetrieverQueryEngine(retriever=retriever)
            
            total_time = time.time() - start_time
            print(f"\n{'='*60}")
            print(f"✓ 内容索引构建完成！")
            print(f"  - 处理文档数：{len(documents)}")
            print(f"  - 总节点数：{len(all_nodes)}")
            print(f"  - 总耗时：{total_time:.2f}秒 ({total_time/60:.2f}分钟)")
            print(f"{'='*60}")
            
        except Exception as e:
            print(f"✗ 内容索引构建失败：{e}")
            import traceback
            traceback.print_exc()
    
    def _read_text_file(self, file_path: Path) -> Optional[Document]:
        """
        读取文本文件（txt, md）
        
        Args:
            file_path: 文件路径
            
        Returns:
            Document 对象
        """
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                
                if not content.strip():
                    return None
                
                return Document(
                    text=content,
                    metadata={
                        'source': str(file_path.relative_to(self.content_dir)),
                        'type': 'text',
                        'file_name': file_path.name
                    }
                )
            except UnicodeDecodeError:
                continue
            except Exception as e:
                print(f"    读取文本文件失败 {file_path.name}: {e}")
                return None
        
        print(f"    无法识别文件编码: {file_path.name}")
        return None
    
    def _read_pdf_file(self, file_path: Path) -> Optional[Document]:
        """
        读取 PDF 文件
        
        Args:
            file_path: PDF 文件路径
            
        Returns:
            Document 对象
        """
        try:
            from llama_index.readers.file import PDFReader
            
            reader = PDFReader()
            pdf_docs = reader.load_data(file=file_path)
            
            content_parts = []
            for doc in pdf_docs:
                content_parts.append(doc.text)
            
            content = "\n\n".join(content_parts)
            
            if not content.strip():
                return None
            
            return Document(
                text=f"[PDF 文件: {file_path.name}]\n{content}",
                metadata={
                    'source': str(file_path.relative_to(self.content_dir)),
                    'type': 'pdf',
                    'file_name': file_path.name
                }
            )
        except Exception as e:
            print(f"    读取 PDF 文件失败 {file_path.name}: {e}")
            return None
    
    def _read_docx_file(self, file_path: Path) -> Optional[Document]:
        """
        读取 Word 文档
        
        Args:
            file_path: Word 文件路径
            
        Returns:
            Document 对象
        """
        try:
            import docx
            
            doc = docx.Document(str(file_path))
            
            content_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)
            
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    content_parts.append("\n[表格]\n" + "\n".join(table_text))
            
            content = "\n\n".join(content_parts)
            
            if not content.strip():
                return None
            
            return Document(
                text=f"[Word 文档: {file_path.name}]\n{content}",
                metadata={
                    'source': str(file_path.relative_to(self.content_dir)),
                    'type': 'word',
                    'file_name': file_path.name
                }
            )
        except ImportError:
            print(f"    警告：python-docx 未安装，无法读取 Word 文件")
            print(f"    请运行: pip install python-docx")
            return None
        except Exception as e:
            print(f"    读取 Word 文件失败 {file_path.name}: {e}")
            return None
    
    def _read_excel_file(self, file_path: Path) -> Optional[Document]:
        """
        读取 Excel 文件
        
        Args:
            file_path: Excel 文件路径
            
        Returns:
            Document 对象
        """
        try:
            import openpyxl
            
            wb = openpyxl.load_workbook(file_path, data_only=True)
            content_parts = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_content = [f"\n## 工作表: {sheet_name}\n"]
                
                max_rows = min(sheet.max_row, 200)
                
                for row_idx in range(1, max_rows + 1):
                    row_data = []
                    for col_idx in range(1, sheet.max_column + 1):
                        cell = sheet.cell(row=row_idx, column=col_idx)
                        cell_value = cell.value
                        if cell_value is not None:
                            row_data.append(str(cell_value))
                    
                    if row_data:
                        sheet_content.append(" | ".join(row_data))
                
                if len(sheet_content) > 1:
                    content_parts.extend(sheet_content)
            
            content = "\n".join(content_parts)
            
            if not content.strip():
                return None
            
            return Document(
                text=f"[Excel 文件: {file_path.name}]\n{content}",
                metadata={
                    'source': str(file_path.relative_to(self.content_dir)),
                    'type': 'excel',
                    'file_name': file_path.name
                }
            )
        except ImportError:
            print(f"    警告：openpyxl 未安装，无法读取 Excel 文件")
            print(f"    请运行: pip install openpyxl")
            return None
        except Exception as e:
            print(f"    读取 Excel 文件失败 {file_path.name}: {e}")
            return None
    
    def _process_image_file(self, image_path: Path) -> Optional[Document]:
        """
        处理图像文件，使用 Qwen-VL 生成描述
        
        Args:
            image_path: 图像文件路径
            
        Returns:
            包含图像描述的 Document 对象
        """
        try:
            import dashscope
            from dashscope import MultiModalConversation
            
            messages = [
                {
                    "role": "user",
                    "content": [
                        {"image": str(image_path)},
                        {"text": "请详细描述这张图片的内容。如果是实验结果图，请说明图表类型、坐标轴含义、数据趋势等。如果是流程图或架构图，请说明各个组件及其关系。如果是其他类型的图片，请尽可能详细地描述其中的文字、数据和关键信息。"}
                    ]
                }
            ]
            
            response = MultiModalConversation.call(
                model='qwen-vl-plus',
                messages=messages
            )
            
            if response.status_code == 200:
                description = response.output.choices[0].message.content
                
                return Document(
                    text=f"[图像文件: {image_path.name}]\n{description}",
                    metadata={
                        'source': str(image_path.relative_to(self.content_dir)),
                        'type': 'image',
                        'file_name': image_path.name
                    }
                )
            else:
                print(f"    Qwen-VL API 调用失败: {response.code} - {response.message}")
                return None
                
        except ImportError:
            print(f"    警告：dashscope 未安装，无法处理图像文件")
            print(f"    请运行: pip install dashscope")
            return None
        except Exception as e:
            print(f"    处理图像文件失败 {image_path.name}: {e}")
            return None
    
    def query(self, query_str: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """
        查询内容索引
        
        Args:
            query_str: 查询字符串
            top_k: 返回结果数量
            
        Returns:
            查询结果列表
        """
        if self.query_engine is None:
            raise RuntimeError("索引未初始化，请先调用 build_index()")
        
        response = self.query_engine.query(query_str)
        
        results = []
        if hasattr(response, 'source_nodes'):
            for node in response.source_nodes:
                results.append({
                    'text': node.node.text,
                    'metadata': node.node.metadata,
                    'score': node.score if hasattr(node, 'score') else None
                })
        
        return results
    
    def retrieve(self, query: str, top_k: int = 5) -> str:
        """
        检索相关内容并返回文本
        
        Args:
            query: 查询字符串
            top_k: 返回结果数量
            
        Returns:
            检索到的内容文本
        """
        results = self.query(query, top_k)
        
        if not results:
            return ""
        
        content_parts = []
        for i, result in enumerate(results, 1):
            source = result['metadata'].get('source', '未知来源')
            text = result['text']
            content_parts.append(f"[参考资料 {i}: {source}]\n{text}")
        
        return "\n\n".join(content_parts)
